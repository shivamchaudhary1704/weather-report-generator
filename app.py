import streamlit as st
from io import BytesIO
from docx import Document
import datetime
from docx.shared import Pt

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.tools import DuckDuckGoSearchRun, WikipediaQueryRun
from langchain_community.utilities import DuckDuckGoSearchAPIWrapper, WikipediaAPIWrapper
from langchain.agents import initialize_agent, Tool, AgentType


api_key = st.secrets["GOOGLE_API_KEY"]
llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash-latest", temperature=0.7)
# ddgo integrated
ddg = DuckDuckGoSearchAPIWrapper()
ddg_tool = DuckDuckGoSearchRun(api_wrapper=ddg)
#added wiki 
wiki = WikipediaAPIWrapper(top_k_results=5)
wiki_tool = WikipediaQueryRun(api_wrapper=wiki)

tools = [
    Tool(
        name="DuckDuckGo",
        func=ddg_tool.run,
        description="Use for general web search and recent info"
    ),
    Tool(
        name="Wikipedia",
        func=wiki_tool.run,
        description="Use for background info, definitions, and summaries"
    )
]

agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=False
)


def generate_report(state: str) -> BytesIO:
    system_prompt = f"""
You are a research assistant. Create a structured research report on: "{state}"
- Use headings: Introduction, Background, Current Trends, Key Data/Stats, Opportunities/Risks, Conclusion
- Do NOT include source URLs inline
- Make it concise but informative
- Use info from Wikipedia and web search (DuckDuckGo)
"""
    try:
        report_text = agent.run(system_prompt)
    except Exception as e:
        st.error("⚠️ API limit reached or an error occurred. Please try again later.")
        st.stop()

    # Create DOCX in memory
    doc = Document()
    doc.add_heading(f"Research Report: {state}", 0)

    for line in report_text.splitlines():
        text = line.strip()
        if not text:
            continue

        # Convert headings to Word-style heading (bigger font)
        if text.startswith("Introduction:"):
            doc.add_heading("Introduction", level=3)
            doc.add_paragraph(text.replace("Introduction:", "").strip())
        elif text.startswith("Background:"):
            doc.add_heading("Background", level=2)
            doc.add_paragraph(text.replace("Background:", "").strip())
        elif text.startswith("Current Trends:"):
            doc.add_heading("Current Trends", level=4)
            doc.add_paragraph(text.replace("Current Trends:", "").strip())
        elif text.startswith("Key Data/Stats:"):
            doc.add_heading("Key Data/Stats", level=1)
            doc.add_paragraph(text.replace("Key Data/Stats:", "").strip())
        elif text.startswith("Opportunities/Risks:"):
            doc.add_heading("Opportunities/Risks", level=1)
            doc.add_paragraph(text.replace("Opportunities/Risks:", "").strip())
        elif text.startswith("Conclusion:"):
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(text.replace("Conclusion:", "").strip())
        else:
            doc.add_paragraph(text)

    doc.add_paragraph("")
    doc.add_paragraph(f"Generated with Gemini · {datetime.date.today().isoformat()}")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="India State Research Report", layout="wide")

# Pink-Blue Gradient + CSS Styling
st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #FFC0CB, #ADD8E6);
    color: #000 !important;
}
.main-title {text-align:center; color:#1E90FF; font-size:42px; font-weight:bold; margin-bottom:10px;}
.subtitle {text-align:center; color:#333; font-size:18px; margin-bottom:30px;}
section[data-testid="stSidebar"] {background: linear-gradient(135deg, #FFB6C1, #87CEFA); color:#000;}
div.stButton > button {background-color:#FF69B4; color:white; font-weight:bold; border-radius:10px; padding:0.6em 1em; transition:all 0.3s ease-in-out;}
div.stButton > button:hover {background-color:#FF1493; transform:scale(1.05);}
div.stDownloadButton > button {background-color:#87CEFA; color:#FF1493; font-weight:bold; border-radius:10px; padding:0.6em 1em; transition:all 0.3s ease-in-out;}
div.stDownloadButton > button:hover {background-color:#00BFFF; color:white; transform:scale(1.05);}
header, footer {visibility: hidden;}
section[data-testid="stHeader"], section[data-testid="stToolbar"] {display:none;}
</style>
""", unsafe_allow_html=True)

# Titles
st.markdown("<div class='main-title'>📄 India State Research Report Generator</div>", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Generate AI-powered research reports for any Indian state 🚀</div>", unsafe_allow_html=True)

# Sidebar - State Selection
st.sidebar.markdown("### 🎯 Select a State")
indian_states = [
    "Andhra Pradesh","Arunachal Pradesh","Assam","Bihar","Chhattisgarh","Goa","Gujarat",
    "Haryana","Himachal Pradesh","Jharkhand","Karnataka","Kerala","Madhya Pradesh",
    "Maharashtra","Manipur","Meghalaya","Mizoram","Nagaland","Odisha","Punjab","Rajasthan",
    "Sikkim","Tamil Nadu","Telangana","Tripura","Uttar Pradesh","Uttarakhand","West Bengal",
    "Delhi","Jammu and Kashmir","Ladakh"
]
selected_state = st.sidebar.selectbox("Choose a state:", indian_states)

# Generate Report Section
if st.button("✨ Generate Report"):
    with st.spinner(f"🔄 Generating report for {selected_state}..."):
        docx_file = generate_report(selected_state)
        st.success("✅ Report successfully generated!")
        # Download button
        st.download_button(
            label="⬇️ Download Full Report (.docx)",
            data=docx_file,
            file_name=f"{selected_state}_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
